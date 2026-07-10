import os
import fitz  # PyMuPDF
import docx

def create_sample_pdf(filepath: str, name: str, email: str, phone: str, skills: list, experience: str, education: str):
    doc = fitz.open()
    page = doc.new_page()
    
    text = f"""
    {name.upper()}
    Email: {email} | Phone: {phone}
    Address: 123 Innovation Drive, Silicon Valley, CA, USA
    
    PROFESSIONAL SUMMARY
    Dynamic Software Engineer with {experience} of experience specializing in building scalable web applications.
    
    EDUCATION
    {education}
    
    SKILLS
    {", ".join(skills)}
    
    WORK EXPERIENCE
    Senior Developer | Tech Innovators Inc. (2022 - Present)
    - Led a team of 4 developers to build a real-time data streaming platform.
    - Designed and implemented microservices using Python and FastAPI.
    
    PROJECTS
    Automated Parsing Engine: Built a pipeline to parse millions of documents using regex and AI.
    
    CERTIFICATIONS
    AWS Certified Solutions Architect
    """
    # Insert text in page
    rect = fitz.Rect(50, 50, 550, 750)
    page.insert_textbox(rect, text, fontsize=11, fontname="Helvetica")
    doc.save(filepath)
    doc.close()
    print(f"Created sample PDF at: {filepath}")

def create_sample_docx(filepath: str, name: str, email: str, phone: str, skills: list, experience: str, education: str):
    doc = docx.Document()
    
    doc.add_heading(name, 0)
    doc.add_paragraph(f"Email: {email} | Phone: {phone} | Location: Austin, TX, USA")
    
    doc.add_heading("Professional Summary", level=1)
    doc.add_paragraph(f"Dedicated professional with {experience} of experience in full-stack engineering and software architecture.")
    
    doc.add_heading("Skills", level=1)
    p = doc.add_paragraph()
    for s in skills:
        p.add_run(f"• {s}  ")
        
    doc.add_heading("Education", level=1)
    doc.add_paragraph(education)
    
    doc.add_heading("Work Experience", level=1)
    doc.add_paragraph("Full Stack Dev | Cloud Systems Ltd (2020 - 2023)")
    doc.add_paragraph("- Designed and deployed React/Node.js web portals for international clients.")
    doc.add_paragraph("- Maintained Postgres databases and optimized query execution times by 40%.")
    
    doc.add_heading("Projects", level=1)
    doc.add_paragraph("HR Screening Tool: Automated candidate screening using LLM intelligence.")
    
    doc.add_heading("Certifications", level=1)
    doc.add_paragraph("Google Cloud Professional Cloud Architect")
    
    doc.save(filepath)
    print(f"Created sample DOCX at: {filepath}")

if __name__ == "__main__":
    os.makedirs("samples", exist_ok=True)
    
    create_sample_pdf(
        "samples/alex_johnson_resume.pdf",
        "Alex Johnson",
        "alex.johnson@example.com",
        "+1-555-0101",
        ["Python", "FastAPI", "SQL", "Docker", "AWS", "Git"],
        "6 years",
        "M.S. in Computer Science - Stanford University (2018)"
    )
    
    create_sample_docx(
        "samples/maria_sanchez_resume.docx",
        "Maria Sanchez",
        "maria.sanchez@example.com",
        "+1-555-0202",
        ["JavaScript", "TypeScript", "React", "Next.js", "Node.js", "Tailwind CSS"],
        "4 years",
        "B.S. in Software Engineering - University of Texas (2020)"
    )
