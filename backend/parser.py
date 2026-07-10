import io
import fitz  # PyMuPDF
import docx

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Extracts text from PDF bytes using PyMuPDF (fitz)
    """
    text = ""
    try:
        # Open PDF from bytes
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        for page in doc:
            text += page.get_text()
        doc.close()
    except Exception as e:
        raise ValueError(f"Failed to parse PDF: {str(e)}")
    return text

def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extracts text from DOCX bytes using python-docx
    """
    text_parts = []
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        
        # Read paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Read tables if any
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_parts.append(" | ".join(row_text))
                    
    except Exception as e:
        raise ValueError(f"Failed to parse DOCX: {str(e)}")
    
    return "\n".join(text_parts)

def extract_text(file_bytes: bytes, filename: str) -> str:
    """
    Detects file type based on extension and extracts text.
    """
    ext = filename.split(".")[-1].lower()
    if ext == "pdf":
        return extract_text_from_pdf(file_bytes)
    elif ext in ["docx", "doc"]:
        return extract_text_from_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: .{ext}. Only PDF and DOCX are supported.")
